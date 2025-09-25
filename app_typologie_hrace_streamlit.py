# -*- coding: utf-8 -*-
# Streamlit app: Typologie hráče – automatický scouting report (CZ)
# Build: v3.4 – běžecká data v textu i grafech, aliasy, debug, přísné doporučení, DOCX

import io
import re
import unicodedata
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import streamlit as st
from matplotlib.colors import LinearSegmentedColormap
from datetime import datetime

# -------------------- Bezpečný import DOCX --------------------
try:
    from docx import Document
    from docx.shared import Inches
    HAVE_DOCX = True
except Exception:
    HAVE_DOCX = False

# -------------------- ZÁKLADNÍ NASTAVENÍ --------------------
st.set_page_config(page_title="Typologie hráče – Scouting report", layout="wide")
st.title("⚽ Typologie hráče – generátor scouting reportu (CZ)")
st.caption("Build: v3.4 – běžecká data v textu i grafech")

# Sidebar
with st.sidebar:
    st.header("⚙️ Nastavení")
    min_minutes = st.number_input("Minimální minuty pro referenční vzorek", min_value=0, value=300, step=30)
    show_percentiles = st.checkbox("Zobrazit percentily (vs. pozice)", value=True)
    tone = st.selectbox("Tón doporučení", ["Přísný", "Neutrální"], index=0)
    st.markdown("### 🏃‍♂️ Běžecká data")
    run_file = st.file_uploader("Nahraj běžecká data (xlsx/csv)", type=["xlsx", "csv"])
    include_runs = st.checkbox("Zahrnout běžecká data do reportu", value=True)
    if st.button("🧹 Clear cache"):
        st.cache_data.clear()
        st.experimental_rerun()
    st.markdown("---")
    st.caption("Tip: Pokud má hráč více řádků (více sezón/zápasů), bere se **první** shoda v datasetu.")

# -------------------- Pomocné funkce --------------------
def safe_float(x):
    try:
        if pd.isna(x): return np.nan
        return float(x)
    except Exception:
        return np.nan

def pct(player_val, league_val):
    if league_val is None or pd.isna(league_val) or league_val == 0 or pd.isna(player_val):
        return np.nan
    return (player_val / league_val) * 100.0

def extract_primary_position(pos_text: str) -> str:
    s = str(pos_text or "").upper()
    m = re.findall(r"[A-Z]{2,3}", s)
    return m[0] if m else s.strip()[:3]

def norm_txt(s: str) -> str:
    s = str(s or "").strip().lower()
    s = unicodedata.normalize("NFKD", s).encode("ascii", "ignore").decode("ascii")
    s = re.sub(r"[^a-z0-9]+", " ", s).strip()
    return s

@st.cache_data(show_spinner=False)
def load_excel(file):
    return pd.read_excel(file)

@st.cache_data(show_spinner=False)
def load_runs_file(file):
    if file.name.lower().endswith(".csv"):
        return pd.read_csv(file)
    return pd.read_excel(file)

def alias_columns(df: pd.DataFrame) -> pd.DataFrame:
    """Mapuj různé názvy sloupců na očekávané aliasy."""
    if df is None or df.empty: return df
    colmap = {}
    # normalizované názvy
    for c in df.columns:
        cn = norm_txt(c)
        # minutes
        if cn in ["minutes", "mins", "minuty", "played_minutes", "time_played"]:
            colmap[c] = "Minutes"
        # total distance
        if cn in ["total distance", "distance total", "td", "distance"]:
            colmap[c] = "Total distance"
        # hir / hsr
        if cn in ["hir distance", "hir", "hsr", "high intensity running", "high speed running"]:
            colmap[c] = "HIR distance"
        # sprints
        if cn in ["sprints", "sprints count", "count sprints", "number of sprints", "sprinty", "sprint count"]:
            colmap[c] = "Sprints count"
        # accelerations
        if cn in ["accelerations", "acc", "accels", "accelerations count"]:
            colmap[c] = "Accelerations"
        # decelerations
        if cn in ["decelerations", "decel", "decels", "decelerations count"]:
            colmap[c] = "Decelerations"
        # per90
        if cn in ["td per 90", "total distance per 90"]:
            colmap[c] = "TD per 90"
        if cn in ["hir per 90", "hsr per 90"]:
            colmap[c] = "HIR per 90"
        if cn in ["sprints per 90", "sprinty per 90"]:
            colmap[c] = "Sprints per 90"
        if cn in ["acc per 90", "accelerations per 90"]:
            colmap[c] = "Acc per 90"
        if cn in ["decel per 90", "decelerations per 90"]:
            colmap[c] = "Decel per 90"
        # top speed
        if cn in ["top speed km h", "top speed", "max speed", "maximal speed", "vmax", "top speed kmh"]:
            colmap[c] = "Top speed (km/h)"
        # repeat sprints per 90
        if cn in ["repeat sprints per 90", "repeated sprints per 90", "rs per 90", "repeated sprint efforts per 90"]:
            colmap[c] = "Repeat sprints per 90"
        # identity map for core keys
        if cn == "player": colmap[c] = "Player"
        if cn == "team": colmap[c] = "Team"
        if cn == "position": colmap[c] = "Position"
    return df.rename(columns=colmap)

def ensure_per90(df, minutes_col="Minutes"):
    df = df.copy()
    m = pd.to_numeric(df.get(minutes_col, np.nan), errors="coerce")
    def mk_per90(src, dst):
        if src in df.columns and dst not in df.columns and m.notna().any():
            df[dst] = pd.to_numeric(df[src], errors="coerce") / (m / 90.0).replace(0, np.nan)
    mk_per90("Total distance", "TD per 90")
    mk_per90("HIR distance", "HIR per 90")
    mk_per90("Sprints count", "Sprints per 90")
    mk_per90("Accelerations", "Acc per 90")
    mk_per90("Decelerations", "Decel per 90")
    return df

def filter_same_position(df: pd.DataFrame, primary_pos: str, min_minutes: int) -> pd.DataFrame:
    if df is None or df.empty: return df
    if "Position" not in df.columns: return df.iloc[0:0].copy()
    pat = rf"\b{re.escape(primary_pos)}\b"
    pos = df["Position"].astype(str).str.upper().str.contains(pat, regex=True, na=False)
    out = df.loc[pos].copy()
    if out.empty:
        pos2 = df["Position"].astype(str).str.upper().str.contains(primary_pos[:2], na=False)
        out = df.loc[pos2].copy()
    if "Minutes played" in out.columns:
        out = out.loc[out["Minutes played"].apply(safe_float) >= min_minutes]
    return out

def pct_rank(series: pd.Series, value: float) -> float:
    s = pd.to_numeric(series, errors="coerce").dropna().sort_values()
    if s.empty or pd.isna(value): return np.nan
    return (s.searchsorted(value, side="right") / len(s)) * 100.0

# -------------------- UI – vstupy --------------------
colA, colB = st.columns([1, 1])
with colA:
    liga_file = st.file_uploader("Nahraj ligový soubor (CZE1.xlsx)", type=["xlsx"])
with colB:
    mode = st.radio("Jak zadáš hráče?", ["Vyberu z ligového souboru", "Nahraju samostatný soubor hráče (xlsx)"])

player_df = None
league_df = None

if liga_file is not None:
    try:
        league_df = load_excel(liga_file)
    except Exception as e:
        st.error(f"Nepodařilo se načíst ligový soubor: {e}")

    if league_df is not None and not league_df.empty:
        if mode == "Vyberu z ligového souboru":
            if "Player" not in league_df.columns:
                st.error("V souboru chybí sloupec 'Player'.")
            else:
                jmena = (
                    league_df.get("Player", pd.Series(dtype=str))
                    .astype(str).dropna().unique().tolist()
                )
                vyber_jmeno = st.selectbox("Vyber hráče ze souboru ligy:", sorted(jmena))
                player_df = league_df.loc[league_df["Player"].astype(str) == vyber_jmeno].copy()
        else:
            player_file = st.file_uploader("Nahraj soubor konkrétního hráče (xlsx)", type=["xlsx"])
            if player_file is not None:
                try:
                    player_df = load_excel(player_file)
                except Exception as e:
                    st.error(f"Nepodařilo se načíst soubor hráče: {e}")

# -------------------- Zpracování --------------------
if league_df is not None and player_df is not None and len(player_df) > 0:
    player_row = player_df.iloc[0]

    jmeno = str(player_row.get("Player", "Neznámý"))
    klub = str(player_row.get("Team", ""))
    pozice_raw = str(player_row.get("Position", ""))
    vek = safe_float(player_row.get("Age", np.nan))
    minuty = safe_float(player_row.get("Minutes played", np.nan))

    primary_pos = extract_primary_position(pozice_raw)
    league_pos = filter_same_position(league_df, primary_pos, min_minutes=min_minutes)
    n_ref = int(len(league_pos))
    league_means = league_pos.mean(numeric_only=True)

    # ---------- BĚŽECKÁ DATA: načtení, aliasy, per90, join ----------
    runs_df = None
    runs_pos = None
    R = {}      # vytažené hodnoty hráče (per90 + top speed)
    RP = {}     # percentily vs. stejná pozice
    matched_rows = 0
    ref_count = 0
    if run_file is not None:
        try:
            runs_df = load_runs_file(run_file)
            runs_df = alias_columns(runs_df)
            runs_df = ensure_per90(runs_df, minutes_col="Minutes")
            # klíče
            for c in ["Player", "Team", "Position"]:
                if c in runs_df.columns:
                    runs_df[c + "_key"] = runs_df[c].apply(norm_txt)
            # filtr minut (pokud existují)
            if "Minutes" in runs_df.columns:
                runs_df = runs_df.loc[pd.to_numeric(runs_df["Minutes"], errors="coerce").fillna(0) >= 90]

            player_key = norm_txt(jmeno)
            team_key   = norm_txt(klub)
            pos_key    = norm_txt(primary_pos)

            # referenční vzorek stejné pozice
            if "Position_key" in runs_df.columns:
                runs_pos = runs_df.loc[runs_df["Position_key"] == pos_key].copy()
            else:
                runs_pos = runs_df.copy()
            ref_count = len(runs_pos)

            # kandidáti
            cand = runs_df
            if {"Player_key", "Team_key"} <= set(runs_df.columns):
                cand = runs_df.loc[(runs_df["Player_key"] == player_key) & (runs_df["Team_key"] == team_key)]
            if cand.empty and "Player_key" in runs_df.columns:
                cand = runs_df.loc[runs_df["Player_key"] == player_key]

            matched_rows = len(cand)
            if matched_rows > 0:
                r = cand.iloc[0]
                for col in ["TD per 90","HIR per 90","Sprints per 90","Acc per 90","Decel per 90","Top speed (km/h)","Repeat sprints per 90"]:
                    if col in runs_df.columns:
                        R[col] = safe_float(r.get(col, np.nan))

                def pct_rank_runs(col):
                    if runs_pos is None or runs_pos.empty or col not in runs_pos.columns:
                        return np.nan
                    return pct_rank(runs_pos[col], R.get(col, np.nan))

                RP = {
                    "td90":  pct_rank_runs("TD per 90"),
                    "hir90": pct_rank_runs("HIR per 90"),
                    "spr90": pct_rank_runs("Sprints per 90"),
                    "acc90": pct_rank_runs("Acc per 90"),
                    "dec90": pct_rank_runs("Decel per 90"),
                    "topv":  pct_rank_runs("Top speed (km/h)"),
                    "rs90":  pct_rank_runs("Repeat sprints per 90"),
                }
        except Exception as e:
            st.warning(f"Běžecká data se nepodařilo zpracovat: {e}")

    # ---------- Stavový panel ----------
    sc1, sc2, sc3, sc4 = st.columns(4)
    sc1.metric("Pozice", primary_pos)
    sc2.metric("Minuty", value=(int(minuty) if not pd.isna(minuty) else "?"))
    sc3.metric("Ref. N (technické)", value=n_ref)
    if run_file is not None:
        sc4.metric("Běžecký vzorek (pozice)", value=(ref_count if ref_count else 0))
        st.caption(f"Match v běžeckých datech: {matched_rows} řádek/řádky. " + ("✅" if matched_rows>0 else "❗ Zkontroluj jméno/klub v souboru."))

    # ---------- Technický RADAR ----------
    st.subheader("Radar – procenta vs. ligový průměr (stejná pozice)")
    radar_map = {
        "Ofenzivní duely vyhrané %": "Offensive duels won, %",
        "Defenzivní duely vyhrané %": "Defensive duels won, %",
        "Hlavičkové souboje vyhrané %": "Aerial duels won, %",
        "Úspěšnost driblinků %": "Successful dribbles, %",
        "Úspěšnost centrů %": "Accurate crosses, %",
        "Góly /90": "Goals per 90",
        "Asistence /90": "Assists per 90",
    }
    radar_labels, radar_vals = [], []
    for lab, col in radar_map.items():
        radar_labels.append(lab)
        radar_vals.append(pct(safe_float(player_row.get(col, np.nan)), safe_float(league_means.get(col, np.nan))))
    vals = np.clip(np.array([0 if pd.isna(v) else v for v in radar_vals]), 0, 150)
    angles = np.linspace(0, 2*np.pi, len(vals), endpoint=False)
    vals_closed = np.concatenate([vals, vals[:1]])
    angles_closed = np.concatenate([angles, angles[:1]])
    fig_radar, ax = plt.subplots(figsize=(3.8, 3.8), subplot_kw=dict(polar=True))
    ax.plot(angles_closed, vals_closed, linewidth=1.6)
    ax.fill(angles_closed, vals_closed, alpha=0.18)
    ax.plot(angles_closed, np.ones_like(vals_closed)*100, linestyle="--", linewidth=1)
    ax.set_xticks(angles); ax.set_xticklabels(radar_labels, fontsize=7)
    ax.set_yticks([50,100,150]); ax.set_yticklabels(["50%","100%","150%"], fontsize=7)
    ax.set_ylim(0,150); ax.grid(alpha=0.35, linewidth=0.6)
    st.pyplot(fig_radar, use_container_width=False); plt.close(fig_radar)

    # ---------- Technická HEATMAP ----------
    st.subheader("Heatmapa – 0–150 % vůči ligovému průměru (stejná pozice)")
    heat_map = {
        "Ofenzivní duely vyhrané %": "Offensive duels won, %",
        "Defenzivní duely vyhrané %": "Defensive duels won, %",
        "Hlavičkové souboje vyhrané %": "Aerial duels won, %",
        "Úspěšnost přihrávek celkem %": "Accurate passes, %",
        "Úspěšnost driblinků %": "Successful dribbles, %",
        "Úspěšnost centrů %": "Accurate crosses, %",
        "Góly /90": "Goals per 90",
        "Asistence /90": "Assists per 90",
    }
    heat_labels, heat_vals = [], []
    for lab, col in heat_map.items():
        heat_labels.append(lab)
        heat_vals.append(pct(safe_float(player_row.get(col, np.nan)), safe_float(league_means.get(col, np.nan))))
    cmap = LinearSegmentedColormap.from_list("r2g", ["#b30000","#ff6b6b","#ffd11a","#b7e1a1","#1e7a1e"])
    hm = np.array([[v if not pd.isna(v) else np.nan] for v in heat_vals], dtype=float)
    hm_plot = np.nan_to_num(hm, nan=100.0)
    fig_hm, ax2 = plt.subplots(figsize=(3.3, 0.26*len(heat_labels)+0.7))
    ax2.imshow(hm_plot, cmap=cmap, vmin=0, vmax=150)
    ax2.set_yticks(range(len(heat_labels))); ax2.set_yticklabels(heat_labels, fontsize=8)
    ax2.set_xticks([])
    for i, v in enumerate(heat_vals):
        txt = "bez dat" if pd.isna(v) else f"{v:.1f}"
        ax2.text(0, i, txt, ha='center', va='center', fontsize=8, fontweight='bold', color='black')
    st.pyplot(fig_hm, use_container_width=False); plt.close(fig_hm)

    # ---------- BĚŽECKÝ RADAR (percentily 0–100) ----------
    if include_runs and run_file is not None and matched_rows > 0 and any(not pd.isna(x) for x in RP.values()):
        st.subheader("🏃‍♂️ Běžecký radar – percentily vs. stejná pozice (0–100)")
        run_labels = []
        run_vals = []
        label_map = [
            ("TD/90",  "td90"),
            ("HIR/90", "hir90"),
            ("Sprinty/90", "spr90"),
            ("Acc/90", "acc90"),
            ("Decel/90", "dec90"),
            ("Top speed", "topv"),
        ]
        for lbl, key in label_map:
            v = RP.get(key, np.nan)
            if not pd.isna(v):
                run_labels.append(lbl); run_vals.append(float(v))
        if run_vals:
            vals2 = np.clip(np.array(run_vals, dtype=float), 0, 100)
            ang = np.linspace(0, 2*np.pi, len(vals2), endpoint=False)
            vals2_c = np.concatenate([vals2, vals2[:1]])
            ang_c = np.concatenate([ang, ang[:1]])
            fig_rr, axr = plt.subplots(figsize=(3.8, 3.8), subplot_kw=dict(polar=True))
            axr.plot(ang_c, vals2_c, linewidth=1.6)
            axr.fill(ang_c, vals2_c, alpha=0.18)
            axr.set_xticks(ang); axr.set_xticklabels(run_labels, fontsize=8)
            axr.set_yticks([25,50,75,100]); axr.set_yticklabels(["25","50","75","100"], fontsize=7)
            axr.set_ylim(0,100); axr.grid(alpha=0.35, linewidth=0.6)
            st.pyplot(fig_rr, use_container_width=False); plt.close(fig_rr)

        # Běžecká heatmapa (percentily)
        st.subheader("🏃‍♂️ Běžecká heatmapa – percentily vs. pozice")
        h_labels = [lbl for lbl,_ in label_map]
        h_vals = [RP.get(key, np.nan) for _,key in label_map]
        h_arr = np.array([[v if not pd.isna(v) else np.nan] for v in h_vals], dtype=float)
        h_plot = np.nan_to_num(h_arr, nan=50.0)
        fig_rhm, axr2 = plt.subplots(figsize=(3.3, 0.26*len(h_labels)+0.7))
        axr2.imshow(h_plot, cmap="viridis", vmin=0, vmax=100)
        axr2.set_yticks(range(len(h_labels))); axr2.set_yticklabels(h_labels, fontsize=8)
        axr2.set_xticks([])
        for i, v in enumerate(h_vals):
            txt = "bez dat" if pd.isna(v) else f"{v:.0f}"
            axr2.text(0, i, txt, ha='center', va='center', fontsize=8, fontweight='bold', color='white')
        st.pyplot(fig_rhm, use_container_width=False); plt.close(fig_rhm)

    # ---------- Percentily tabulka (volitelné) ----------
    if show_percentiles:
        st.subheader("Percentily hráče vs. vzorek stejné pozice")
        rows = []
        for lab, col in heat_map.items():
            val = safe_float(player_row.get(col, np.nan))
            prc = pct_rank(league_pos[col] if col in league_pos.columns else pd.Series(dtype=float), val)
            rows.append({"Metrika": lab, "Hodnota": val, "Percentil": (None if pd.isna(prc) else round(prc, 1))})
        df_prc = pd.DataFrame(rows)
        st.dataframe(df_prc, use_container_width=True)

    # -------------------- DLOUHÝ NARATIV – TYPOLOGIE --------------------
    st.subheader("Scouting report – souvislý text (narativ)")

    # helpery pro text
    def pct_rank_safe(col_name):
        if col_name not in league_pos.columns: return np.nan
        return pct_rank(league_pos[col_name], safe_float(player_row.get(col_name, np.nan)))

    def bucket(p):
        if pd.isna(p): return "bez_dat"
        if p < 25:     return "slabe"
        if p < 50:     return "podprum"
        if p < 75:     return "nadprum"
        return "elite"

    def fmt_percentil(p):
        return "bez dat" if pd.isna(p) else f"{p:.0f}. percentil"

    def fmt_num(v, nd=2):
        v2 = safe_float(v); return "bez dat" if pd.isna(v2) else f"{v2:.{nd}f}"

    P = {
        "drib":    pct_rank_safe("Successful dribbles, %"),
        "cross":   pct_rank_safe("Accurate crosses, %"),
        "aerial":  pct_rank_safe("Aerial duels won, %"),
        "offd":    pct_rank_safe("Offensive duels won, %"),
        "defd":    pct_rank_safe("Defensive duels won, %"),
        "passacc": pct_rank_safe("Accurate passes, %"),
        "shots90": pct_rank_safe("Shots per 90"),
        "touch90": pct_rank_safe("Touches in box per 90"),
        "keyp90":  pct_rank_safe("Key passes per 90"),
        "g90":     pct_rank_safe("Goals per 90"),
        "a90":     pct_rank_safe("Assists per 90"),
    }

    # archetyp + běžecký hint
    def archetype_text(pos, P):
        if pos == "CB":
            core = "silový stoper" if bucket(P["aerial"]) in ["nadprum","elite"] else "poziční stoper"
            ball = "slušná první rozehrávka" if bucket(P["passacc"]) in ["nadprum","elite"] else "jednoduchá rozehrávka"
            trio = "vhodný i do trojice stoperů" if bucket(P["aerial"]) in ["nadprum","elite"] else "vhodnější do kompaktní dvojice"
            return core + ", " + ball + ", " + trio
        if pos in ["RB","LB","RWB","LWB"]:
            if bucket(P["cross"]) in ["nadprum","elite"]:
                return "ofenzivní bek/wingback s doručováním z kraje"
            return "bek orientovaný na defenzivní stabilitu a krytí prostoru"
        if pos in ["LW","RW","LWF","RWF"]:
            if bucket(P["drib"]) in ["nadprum","elite"] and bucket(P["keyp90"]) in ["nadprum","elite"]:
                return "křídlo-playmaker do 1v1 a poslední přihrávky"
            if bucket(P["g90"]) in ["nadprum","elite"] and bucket(P["touch90"]) in ["nadprum","elite"]:
                return "přímočaré křídlo/AMF s náběhy do boxu"
            return "křídlo do přechodu a otevřeného prostoru"
        if pos in ["CF","ST","CF9"]:
            if bucket(P["g90"]) in ["nadprum","elite"]:
                return "boxový zakončovatel"
            if bucket(P["aerial"]) in ["nadprum","elite"]:
                return "target útočník do hry do těla"
            return "spojka pro kombinaci"
        return "univerzální profil"

    def running_hint(RP):
        parts = []
        if RP:
            if not pd.isna(RP.get("topv", np.nan)):
                parts.append("maximální rychlost " + fmt_percentil(RP["topv"]))
            if not pd.isna(RP.get("spr90", np.nan)):
                parts.append("objem sprintů " + ("nadprůměr" if RP["spr90"] >= 60 else "ligový standard" if RP["spr90"] >= 40 else "podprůměr"))
            if not pd.isna(RP.get("hir90", np.nan)):
                parts.append("high-intensity running " + ("nadprůměr" if RP["hir90"] >= 60 else "ligový standard" if RP["hir90"] >= 40 else "podprůměr"))
        return ", ".join(parts) if parts else "rychlostně spíše průměrný"

    arche = archetype_text(primary_pos, P)
    phys_hint = running_hint(RP)

    goals90 = fmt_num(player_row.get("Goals per 90", np.nan))
    shots90 = fmt_num(player_row.get("Shots per 90", np.nan))
    assists90 = fmt_num(player_row.get("Assists per 90", np.nan))
    keyp90_v = fmt_num(player_row.get("Key passes per 90", np.nan))
    touch90_v = fmt_num(player_row.get("Touches in box per 90", np.nan))

    minutes_ok = (not pd.isna(minuty)) and (minuty >= max(600, min_minutes))
    intro = (
        f"{jmeno} ({int(vek) if not pd.isna(vek) else 'věk ?'}, {klub}) je {arche}. "
        + (f"Běžecky: {phys_hint}. " if include_runs and run_file is not None else "")
        + f"V této sezóně odehrál {int(minuty) if not pd.isna(minuty) else '?'} minut – vzorek je "
        + ("dostatečný" if minutes_ok else "omezený")
        + f". Porovnáváno se vzorkem stejné pozice ({primary_pos}, N={n_ref})."
    )

    # detaily běhu v závorce
    run_detail = ""
    if include_runs and run_file is not None and matched_rows > 0:
        parts = []
        if "Top speed (km/h)" in R and not pd.isna(R.get("Top speed (km/h)")):
            parts.append(f"top speed: {R['Top speed (km/h)']:.1f} km/h")
        for name, key in [("TD/90","TD per 90"),("HIR/90","HIR per 90"),("sprinty/90","Sprints per 90")]:
            if key in R and not pd.isna(R.get(key)):
                parts.append(f"{name}: {R[key]:.1f}")
        if parts:
            intro += " (" + "; ".join(parts) + ")."

    # produkce + technika
    if bucket(P["g90"]) in ["slabe","podprum"] and bucket(P["a90"]) in ["slabe","podprum"]:
        prod_clause = "finální výstup je slabý vzhledem k objemu"
    elif bucket(P["g90"]) in ["nadprum","elite"] or bucket(P["a90"]) in ["nadprum","elite"]:
        prod_clause = "finální výstup drží ligový standard nebo nad ním"
    else:
        prod_clause = "finální výstup kolísá kolem průměru"

    tech_bits = []
    if bucket(P["drib"]) in ["nadprum","elite"]: tech_bits.append("silný v driblinku")
    if bucket(P["cross"]) in ["nadprum","elite"]: tech_bits.append("dovede nadprůměrně centrovat")
    pass_note = "celková přesnost přihrávek je v normě" if bucket(P["passacc"]) not in ["slabe"] else "přesnost přihrávek je slabší pod tlakem"

    prod = (
        "Do akcí se dostává pravidelně; " + prod_clause + ". " +
        "Technicky " + (", ".join(tech_bits) if tech_bits else "volí spíše bezpečná řešení") + "; " + pass_note + ". " +
        "Tvorba pro spoluhráče je " + ("nadprůměrná" if bucket(P["keyp90"]) in ["nadprum","elite"] else "spíše pod průměrem") +
        f" (key passes/90: {keyp90_v}). " +
        "Zakončení vyžaduje " + ("lépe volit první dotyk v boxu" if bucket(P["g90"]) in ["slabe","podprum"] else "přenést současné vzorce do vyšší zátěže") +
        f" (góly/90: {goals90}, střely/90: {shots90}; doteky v boxu/90: {touch90_v})."
    )

    # souboje
    duel_txt = []
    if primary_pos in ["LW","RW","LWF","RWF","AMF","CF","ST","CF9"]:
        duel_txt.append("v ofenzivních duelech je nadprůměrný a objemově aktivní" if bucket(P["offd"]) in ["nadprum","elite"] else "ofenzivní duely drží spíše průměr")
    duel_txt.append("defenzivně je spolehlivý v 1v1 a načasování" if bucket(P["defd"]) in ["nadprum","elite"] else "defenzivně je spíše poziční")
    if bucket(P["aerial"]) in ["slabe","podprum"]:
        duel_txt.append("ve vzduchu pod průměrem; target role mu nesedí")
    elif bucket(P["aerial"]) in ["nadprum","elite"]:
        duel_txt.append("ve vzduchu nad průměrem; kryje zadní tyč i standardky")
    duels = "V soubojové činnosti " + ", ".join(duel_txt) + "."

    # herní styl a přenositelnost
    fit_bits = []
    if primary_pos == "CB":
        fit_bits.append("vhodný do struktur s trojicí stoperů a zajištěním prostoru")
        if bucket(P["passacc"]) in ["nadprum","elite"]:
            fit_bits.append("v dvojici zvládne první rozehrávku vedle mobilnějšího partnera")
    elif primary_pos in ["RB","LB","RWB","LWB"]:
        fit_bits.append("sedí do přechodové hry a vysokého postavení krajů")
    elif primary_pos in ["LW","RW","LWF","RWF","AMF"]:
        fit_bits.append("silný v otevřeném prostoru a proti nekompaktním obranám")
        fit_bits.append("proti hlubokému bloku vliv klesá, pokud nemá stabilní rozhodující moment")
    elif primary_pos in ["CF","ST","CF9"]:
        fit_bits.append("uplatní se v boxových vzorcích (cutback, zadní tyč) a řízeném presinku")
    else:
        fit_bits.append("role dle plánu, důležitá je kompaktnost mezi liniemi")
    style = "Herně mu sedí " + ", ".join(fit_bits) + "."

    # přísné doporučení
    def strict_recommendation(pos, P, tone):
        upper = (bucket(P["g90"]) in ["elite","nadprum"] or bucket(P["a90"]) in ["elite","nadprum"] or
                 (pos=="CB" and (bucket(P["aerial"]) in ["elite","nadprum"] and bucket(P["passacc"]) in ["nadprum","elite"])))
        mid = (bucket(P["defd"]) in ["nadprum","elite"] or bucket(P["aerial"]) in ["nadprum","elite"] or
               bucket(P["cross"]) in ["nadprum","elite"] or bucket(P["drib"]) in ["nadprum","elite"])
        if tone == "Přísný":
            if not upper and not mid:
                return "Do top trojky jej nyní nedoporučuji; smysl dává dolní až střední patro tabulky jako šířka kádru."
            if upper:
                return "Vhodný pro ambiciózní horní polovinu tabulky; do dominantního prostředí pouze s jasnou rolí a ochranou slabin."
            return "Použitelný pro střed tabulky; do top projektů jen pod konkrétní herní plán, jinak nedoporučuji."
        return "Reálně využitelný pro střed až horní střed tabulky; do topu po potvrzení konzistence ve finální třetině."

    rec = "Doporučení: " + strict_recommendation(primary_pos, P, tone)

    paragraphs = [intro, prod, duels, style, rec]
    st.write("\n\n".join(paragraphs))

    # -------------------- Export reportu (DOCX) --------------------
    st.subheader("Export reportu (DOCX)")
    if HAVE_DOCX:
        # re-render technického radaru
        buf_radar = io.BytesIO()
        fig_r, ax_r = plt.subplots(figsize=(3.8, 3.8), subplot_kw=dict(polar=True))
        ax_r.plot(angles_closed, vals_closed, linewidth=1.6); ax_r.fill(angles_closed, vals_closed, alpha=0.18)
        ax_r.plot(angles_closed, np.ones_like(vals_closed)*100, linestyle="--", linewidth=1)
        ax_r.set_xticks(angles); ax_r.set_xticklabels(radar_labels, fontsize=7)
        ax_r.set_yticks([50,100,150]); ax_r.set_yticklabels(["50%","100%","150%"], fontsize=7)
        ax_r.set_ylim(0,150)
        fig_r.savefig(buf_radar, format='png', dpi=200, bbox_inches='tight'); plt.close(fig_r); buf_radar.seek(0)

        # re-render technické heatmapy
        buf_heat = io.BytesIO()
        fig_h, ax_h = plt.subplots(figsize=(3.3, 0.26*len(heat_labels)+0.7))
        ax_h.imshow(hm_plot, cmap=cmap, vmin=0, vmax=150)
        ax_h.set_yticks(range(len(heat_labels))); ax_h.set_yticklabels(heat_labels, fontsize=8)
        ax_h.set_xticks([]); 
        for i, v in enumerate(heat_vals):
            txt = "bez dat" if pd.isna(v) else f"{v:.1f}"
            ax_h.text(0, i, txt, ha='center', va='center', fontsize=8, fontweight='bold', color='black')
        fig_h.savefig(buf_heat, format='png', dpi=200, bbox_inches='tight'); plt.close(fig_h); buf_heat.seek(0)

        # volitelně běžecký radar
        buf_run = None
        if include_runs and run_file is not None and matched_rows > 0 and any(not pd.isna(x) for x in RP.values()):
            buf_run = io.BytesIO()
            lbls = []; valsr = []
            for lbl,key in [("TD/90","td90"),("HIR/90","hir90"),("Sprinty/90","spr90"),("Acc/90","acc90"),("Decel/90","dec90"),("Top speed","topv")]:
                v = RP.get(key, np.nan)
                if not pd.isna(v): lbls.append(lbl); valsr.append(float(v))
            if valsr:
                valsr = np.clip(np.array(valsr, dtype=float), 0, 100)
                ang = np.linspace(0, 2*np.pi, len(valsr), endpoint=False)
                valsr_c = np.concatenate([valsr, valsr[:1]]); ang_c = np.concatenate([ang, ang[:1]])
                fig_rb, ax_rb = plt.subplots(figsize=(3.6, 3.6), subplot_kw=dict(polar=True))
                ax_rb.plot(ang_c, valsr_c, linewidth=1.6); ax_rb.fill(ang_c, valsr_c, alpha=0.18)
                ax_rb.set_xticks(ang); ax_rb.set_xticklabels(lbls, fontsize=7)
                ax_rb.set_yticks([25,50,75,100]); ax_rb.set_yticklabels(["25","50","75","100"], fontsize=7)
                ax_rb.set_ylim(0,100)
                fig_rb.savefig(buf_run, format='png', dpi=200, bbox_inches='tight'); plt.close(fig_rb); buf_run.seek(0)

        doc = Document()
        doc.add_heading(f"Typologie hráče – {jmeno}", level=0)
        meta = doc.add_paragraph()
        meta.add_run(f"Klub: {klub} | Pozice: {pozice_raw}\n").bold = True
        meta.add_run(f"Referenční vzorek: stejná pozice {primary_pos}, N={n_ref}\n")
        meta.add_run(f"Minuty: {int(minuty) if not pd.isna(minuty) else '?'} | Věk: {int(vek) if not pd.isna(vek) else '?'}\n")
        doc.add_paragraph(datetime.now().strftime("Vygenerováno: %d.%m.%Y %H:%M"))

        doc.add_heading("Narativní scouting report", level=1)
        for p in paragraphs: doc.add_paragraph(p)

        doc.add_heading("Technický radar (% vs. liga – stejná pozice)", level=1)
        doc.add_picture(buf_radar, width=Inches(3.2))

        doc.add_heading("Technická heatmapa (0–150 % – stejná pozice)", level=1)
        doc.add_picture(buf_heat, width=Inches(3.0))

        if buf_run is not None:
            doc.add_heading("Běžecký radar (percentily vs. pozice)", level=1)
            doc.add_picture(buf_run, width=Inches(3.0))

        out = io.BytesIO(); doc.save(out); out.seek(0)
        st.download_button(
            label="⬇️ Stáhnout scouting report (.docx)",
            data=out,
            file_name=f"Typologie_{jmeno.replace(' ','_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.warning("DOCX export není dostupný – chybí python-docx. Přidej ho do requirements.txt a redeployni appku.")

else:
    st.info("Nahraj ligový soubor a vyber/nahraj hráče – pak ti vygeneruju vizuály a stažitelný report.")



