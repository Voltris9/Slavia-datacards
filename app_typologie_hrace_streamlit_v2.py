# -*- coding: utf-8 -*-
# Streamlit app: Typologie hráče – automatický scouting report (CZ)
# Build: v2.4 – detailní report s percentily, archetypy a taktickým fitem

import io
import re
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import streamlit as st
from matplotlib.colors import LinearSegmentedColormap
from datetime import datetime

# -------------------- Bezpečný import DOCX --------------------
try:
    from docx import Document
    from docx.shared import Inches
    HAVE_DOCX = True
except Exception:
    HAVE_DOCX = False

# -------------------- ZÁKLADNÍ NASTAVENÍ --------------------
st.set_page_config(page_title="Typologie hráče – Scouting report", layout="wide")
st.title("⚽ Typologie hráče – generátor scouting reportu (CZ)")
st.caption("Build: v2.4 – detailní report s percentily, archetypy a taktickým fitem")

# SideBar – globální volby
with st.sidebar:
    st.header("⚙️ Nastavení")
    low_band = st.slider("Hranice 'podprůměr' (LOW)", 50, 95, 70, 1)
    high_band = st.slider("Hranice 'nadprůměr' (HIGH)", 105, 180, 130, 1)
    min_minutes = st.number_input("Minimální minuty pro referenční vzorek", min_value=0, value=300, step=30)
    show_percentiles = st.checkbox("Zobrazit percentily (vs. pozice)", value=True)
    if st.button("🧹 Clear cache"):
        st.cache_data.clear()
        st.experimental_rerun()
    st.markdown("---")
    st.caption("Tip: Pokud má hráč více řádků (více sezón/zápasů), bere se **první** shoda v datasetu.")

# -------------------- Pomocné funkce --------------------
def safe_float(x):
    try:
        if pd.isna(x):
            return np.nan
        return float(x)
    except Exception:
        return np.nan

# Poměr hráč / ligový průměr v %
def pct(player_val, league_val):
    if league_val is None or pd.isna(league_val) or league_val == 0 or pd.isna(player_val):
        return np.nan
    return (player_val / league_val) * 100.0

# Slovní pásma (jen pro rychlé označení)
def band(v, low=70, high=130):
    if pd.isna(v):
        return "bez dat"
    if v < low:
        return "podprůměr"
    if v > high:
        return "nadprůměr"
    return "ligový standard"

def band_adj(word):
    return {
        "podprůměr": "podprůměrný",
        "nadprůměr": "nadprůměrný",
        "ligový standard": "ligový standard",
        "bez dat": "bez dat"
    }.get(word, word)

# Primární pozice – první dvou/třípísmenný token
def extract_primary_position(pos_text: str) -> str:
    s = str(pos_text or "").upper()
    m = re.findall(r"[A-Z]{2,3}", s)
    return m[0] if m else s.strip()[:3]

# -------------------- CACHE I/O --------------------
@st.cache_data(show_spinner=False)
def load_excel(file):
    return pd.read_excel(file)

# Filtrovat ligu jen na stejnou pozici (a minuty)
def filter_same_position(df: pd.DataFrame, primary_pos: str, min_minutes: int) -> pd.DataFrame:
    if df is None or df.empty:
        return df
    if "Position" not in df.columns:
        return df.iloc[0:0].copy()

    pat = rf"\b{re.escape(primary_pos)}\b"
    pos = df["Position"].astype(str).str.upper().str.contains(pat, regex=True, na=False)
    out = df.loc[pos].copy()

    # záloha dle prefixu, kdyby nic
    if out.empty:
        pos2 = df["Position"].astype(str).str.upper().str.contains(primary_pos[:2], na=False)
        out = df.loc[pos2].copy()

    # filtr minut (pokud sloupec existuje)
    if "Minutes played" in out.columns:
        out = out.loc[out["Minutes played"].apply(safe_float) >= min_minutes]

    return out

# Percentily pro srovnání v rámci pozice
def pct_rank(series: pd.Series, value: float) -> float:
    s = pd.to_numeric(series, errors="coerce").dropna().sort_values()
    if s.empty or pd.isna(value):
        return np.nan
    return (s.searchsorted(value, side="right") / len(s)) * 100.0

# -------------------- UI – vstupy --------------------
colA, colB = st.columns([1, 1])
with colA:
    liga_file = st.file_uploader("Nahraj ligový soubor (CZE1.xlsx)", type=["xlsx"])
with colB:
    mode = st.radio("Jak zadáš hráče?", ["Vyberu z ligového souboru", "Nahraju samostatný soubor hráče (xlsx)"])

player_df = None
league_df = None

if liga_file is not None:
    try:
        league_df = load_excel(liga_file)
    except Exception as e:
        st.error(f"Nepodařilo se načíst ligový soubor: {e}")

    if league_df is not None and not league_df.empty:
        if mode == "Vyberu z ligového souboru":
            if "Player" not in league_df.columns:
                st.error("V souboru chybí sloupec 'Player'.")
            else:
                jmena = (
                    league_df.get("Player", pd.Series(dtype=str))
                    .astype(str)
                    .dropna()
                    .unique()
                    .tolist()
                )
                vyber_jmeno = st.selectbox("Vyber hráče ze souboru ligy:", sorted(jmena))
                player_df = league_df.loc[league_df["Player"].astype(str) == vyber_jmeno].copy()
        else:
            player_file = st.file_uploader("Nahraj soubor konkrétního hráče (xlsx)", type=["xlsx"])
            if player_file is not None:
                try:
                    player_df = load_excel(player_file)
                except Exception as e:
                    st.error(f"Nepodařilo se načíst soubor hráče: {e}")

# -------------------- Zpracování --------------------
required_cols = [
    "Player", "Team", "Position", "Age", "Minutes played",
]

if league_df is not None and player_df is not None and len(player_df) > 0:
    # Validace sloupců
    missing = [c for c in required_cols if c not in pd.concat([league_df, player_df], axis=0).columns]
    if missing:
        st.warning("Chybí následující klíčové sloupce: " + ", ".join(missing))

    player_row = player_df.iloc[0]

    # Základní metadata
    jmeno = str(player_row.get("Player", "Neznámý"))
    klub = str(player_row.get("Team", ""))
    pozice_raw = str(player_row.get("Position", ""))
    vek = safe_float(player_row.get("Age", np.nan))
    minuty = safe_float(player_row.get("Minutes played", np.nan))

    # Stejná pozice – referenční vzorek
    primary_pos = extract_primary_position(pozice_raw)
    league_pos = filter_same_position(league_df, primary_pos, min_minutes=min_minutes)
    n_ref = int(len(league_pos))
    league_means = league_pos.mean(numeric_only=True)

    # Metriky pro RADAR a HEATMAPU (v % vůči průměru na stejné pozici)
    radar_map = {
        "Ofenzivní duely vyhrané %": "Offensive duels won, %",
        "Defenzivní duely vyhrané %": "Defensive duels won, %",
        "Hlavičkové souboje vyhrané %": "Aerial duels won, %",
        "Úspěšnost driblinků %": "Successful dribbles, %",
        "Úspěšnost centrů %": "Accurate crosses, %",
        "Góly /90": "Goals per 90",
        "Asistence /90": "Assists per 90",
    }

    heat_map = {
        "Ofenzivní duely vyhrané %": "Offensive duels won, %",
        "Defenzivní duely vyhrané %": "Defensive duels won, %",
        "Hlavičkové souboje vyhrané %": "Aerial duels won, %",
        "Úspěšnost přihrávek celkem %": "Accurate passes, %",
        "Úspěšnost driblinků %": "Successful dribbles, %",
        "Úspěšnost centrů %": "Accurate crosses, %",
        "Góly /90": "Goals per 90",
        "Asistence /90": "Assists per 90",
    }

    # Další metriky pro text (aktivita/kreativita)
    shots90   = safe_float(player_row.get("Shots per 90", np.nan))
    shots90_L = safe_float(league_means.get("Shots per 90", np.nan))
    touch90   = safe_float(player_row.get("Touches in box per 90", np.nan))
    touch90_L = safe_float(league_means.get("Touches in box per 90", np.nan))
    keyp90    = safe_float(player_row.get("Key passes per 90", np.nan))
    keyp90_L  = safe_float(league_means.get("Key passes per 90", np.nan))

    # Výpočty % vs liga (stejná pozice)
    radar_labels, radar_vals = [], []
    for lab, col in radar_map.items():
        radar_labels.append(lab)
        radar_vals.append(pct(safe_float(player_row.get(col, np.nan)), safe_float(league_means.get(col, np.nan))))

    heat_labels, heat_vals = [], []
    for lab, col in heat_map.items():
        heat_labels.append(lab)
        heat_vals.append(pct(safe_float(player_row.get(col, np.nan)), safe_float(league_means.get(col, np.nan))))

    # -------------------- Info box --------------------
    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Pozice", primary_pos)
    col2.metric("Minuty", value=(int(minuty) if not pd.isna(minuty) else "?"))
    col3.metric("Referenční N", value=n_ref)
    col4.metric("Věk", value=(int(vek) if not pd.isna(vek) else "?"))

    # -------------------- Vizualizace: RADAR --------------------
    st.subheader("Radar – procenta vs. ligový průměr (stejná pozice)")
    vals = np.array([v if not pd.isna(v) else 0 for v in radar_vals])
    vals = np.clip(vals, 0, 150)
    angles = np.linspace(0, 2*np.pi, len(vals), endpoint=False)
    vals_closed = np.concatenate([vals, vals[:1]])
    angles_closed = np.concatenate([angles, angles[:1]])

    fig_radar, ax = plt.subplots(figsize=(3.8, 3.8), subplot_kw=dict(polar=True))
    ax.plot(angles_closed, vals_closed, linewidth=1.6, label=jmeno)
    ax.fill(angles_closed, vals_closed, alpha=0.18)
    ax.plot(angles_closed, np.ones_like(vals_closed)*100, linestyle="--", linewidth=1, label=f"Liga = 100% ({primary_pos}, N={n_ref})")
    ax.set_xticks(angles)
    ax.set_xticklabels(radar_labels, fontsize=7)
    ax.set_yticks([50, 100, 150])
    ax.set_yticklabels(["50%", "100%", "150%"], fontsize=7)
    ax.set_ylim(0, 150)
    ax.grid(alpha=0.35, linewidth=0.6)
    ax.set_title(f"{jmeno} – radar (% vs. liga, {primary_pos})", fontsize=9)
    ax.legend(loc="upper right", bbox_to_anchor=(1.22, 1.05), prop={'size': 7})
    st.pyplot(fig_radar, use_container_width=False)
    plt.close(fig_radar)

    # -------------------- Vizualizace: HEATMAP --------------------
    st.subheader("Heatmapa – 0–150 % vůči ligovému průměru (stejná pozice)")
    cmap = LinearSegmentedColormap.from_list("r2g", ["#b30000", "#ff6b6b", "#ffd11a", "#b7e1a1", "#1e7a1e"])

    hm = np.array([[v if not pd.isna(v) else np.nan] for v in heat_vals], dtype=float)
    hm_plot = np.nan_to_num(hm, nan=100.0)

    fig_hm, ax2 = plt.subplots(figsize=(3.3, 0.26*len(heat_labels)+0.7))
    ax2.imshow(hm_plot, cmap=cmap, vmin=0, vmax=150)
    ax2.set_yticks(range(len(heat_labels)))
    ax2.set_yticklabels(heat_labels, fontsize=8)
    ax2.set_xticks([])
    for i, v in enumerate(heat_vals):
        txt = "bez dat" if pd.isna(v) else f"{v:.1f}"
        ax2.text(0, i, txt, ha='center', va='center', fontsize=8, fontweight='bold', color='black')
    ax2.set_title(f"{jmeno} – komplexní činnosti (0–150 %, {primary_pos}, N={n_ref})", fontsize=9)
    st.pyplot(fig_hm, use_container_width=False)
    plt.close(fig_hm)

    # -------------------- Percentily tabulka (volitelné) --------------------
    if show_percentiles:
        st.subheader("Percentily hráče vs. vzorek stejné pozice")
        rows = []
        for lab, col in heat_map.items():
            val = safe_float(player_row.get(col, np.nan))
            prc = pct_rank(league_pos[col] if col in league_pos.columns else pd.Series(dtype=float), val)
            rows.append({"Metrika": lab, "Hodnota": val, "Percentil": (None if pd.isna(prc) else round(prc, 1))})
        df_prc = pd.DataFrame(rows)
        st.dataframe(df_prc, use_container_width=True)

   # -------------------- Textový report (narrativní/strukturovaný) --------------------
st.subheader("Scouting report – souvislý text")

# Ovládání stylu a tónu
col_style, col_tone = st.columns([1,1])
with col_style:
    report_style = st.selectbox("Styl reportu", ["Narativní typologie (doporučeno)", "Strukturovaný (původní)"], index=0)
with col_tone:
    tone = st.selectbox("Tón doporučení", ["Přísný", "Neutrální"], index=0)

# helpery zůstávají stejné
def pct_rank_safe(col_name):
    if col_name not in league_pos.columns:
        return np.nan
    return pct_rank(league_pos[col_name], safe_float(player_row.get(col_name, np.nan)))

def bucket(p):
    if pd.isna(p): return "bez_dat"
    if p < 25:     return "slabe"
    if p < 50:     return "podprum"
    if p < 75:     return "nadprum"
    return "elite"

def fmt_percentil(p):
    return "bez dat" if pd.isna(p) else f"{p:.0f}. percentil"

def val(v, nd=2):
    v2 = safe_float(v)
    return "bez dat" if pd.isna(v2) else f"{v2:.{nd}f}"

# percentily pro rozhodování
P = {
    "drib":    pct_rank_safe("Successful dribbles, %"),
    "cross":   pct_rank_safe("Accurate crosses, %"),
    "aerial":  pct_rank_safe("Aerial duels won, %"),
    "offd":    pct_rank_safe("Offensive duels won, %"),
    "defd":    pct_rank_safe("Defensive duels won, %"),
    "passacc": pct_rank_safe("Accurate passes, %"),
    "shots90": pct_rank_safe("Shots per 90"),
    "touch90": pct_rank_safe("Touches in box per 90"),
    "keyp90":  pct_rank_safe("Key passes per 90"),
    "g90":     pct_rank_safe("Goals per 90"),
    "a90":     pct_rank_safe("Assists per 90"),
}

# --- Funkce: slovní archetyp podle pozice a dat (stručný tag + věta) ---
def archetype_sentence(pos, P):
    tags = []
    sent = ""
    if pos in ["CB"]:
        if bucket(P["aerial"]) in ["elite","nadprum"] and bucket(P["defd"]) in ["nadprum","elite"]:
            tags.append("silový stoper")
        if bucket(P["passacc"]) in ["nadprum","elite"]:
            tags.append("klid v první rozehrávce")
        if bucket(P["aerial"]) in ["nadprum","elite"] and bucket(P["passacc"]) not in ["nadprum","elite"]:
            sent = "Profilově jde o silového stopera do struktur s trojicí vzadu; ve vzduchu přináší stabilitu, na míči hraje jednoduše."
        elif bucket(P["aerial"]) in ["nadprum","elite"] and bucket(P["passacc"]) in ["nadprum","elite"]:
            sent = "Typologicky stoper s dominancí ve vzduchu a slušnou první rozehrávkou; může hrát ve dvojici i v trojici."
        else:
            sent = "Stoper se zaměřením na obranu prostoru a hlavičkové situace; rozehrávku drží spíše bezpečnou."
    elif pos in ["RB","LB","RWB","LWB"]:
        if bucket(P["cross"]) in ["nadprum","elite"]:
            tags.append("ofenzivní bek/wingback – doručování")
        if bucket(P["defd"]) in ["nadprum","elite"]:
            tags.append("spolehlivý 1v1 v defenzivě")
        sent = "Bek vhodný do přechodové hry a vysokého postavení; doručování z kraje je nadprůměrné." \
               if bucket(P["cross"]) in ["nadprum","elite"] else \
               "Bek se zaměřením na defenzivní stabilitu; lepší v 1v1 než v tvorbě z hloubky."
    elif pos in ["LW","RW","LWF","RWF"]:
        if bucket(P["drib"]) in ["nadprum","elite"] and bucket(P["keyp90"]) in ["nadprum","elite"]:
            tags.append("křídlo-playmaker")
            sent = "Křídlo se schopností přechodu 1v1 a poslední přihrávky; hrozí z halfspace."
        elif bucket(P["g90"]) in ["nadprum","elite"] and bucket(P["touch90"]) in ["nadprum","elite"]:
            tags.append("křídlo-finisher")
            sent = "Přímočaré křídlo s pohybem do boxu a zakončením na zadní tyči."
        else:
            tags.append("křídlo-transition")
            sent = "Křídlo vhodné do otevřeného prostoru; v bloku je produkce proměnlivá."
    elif pos in ["CF","ST","CF9"]:
        if bucket(P["aerial"]) in ["nadprum","elite"]:
            tags.append("9 – target")
            sent = "Útočník vhodný pro kombinační ukotvení a hru do těla; využitelný na vysoké míče."
        elif bucket(P["g90"]) in ["nadprum","elite"]:
            tags.append("9 – finisher")
            sent = "Boxový zakončovatel se smyslem pro načasování v šestnáctce."
        else:
            tags.append("9 – spojka")
            sent = "Útočník pro spojení hry; finální výstup je spíše průměrný."
    else:
        tags.append("univerzální profil")
        sent = "Univerzální středový hráč; přidaná hodnota se odvíjí od kontextu herního plánu."
    return ", ".join(tags), sent

# --- Funkce: přísné doporučení (konzervativní) ---
def strict_recommendation(pos, P, tone):
    upper = (bucket(P["g90"]) in ["elite","nadprum"] or bucket(P["keyp90"]) in ["elite","nadprum"] or
             (pos=="CB" and (bucket(P["aerial"]) in ["elite","nadprum"] and bucket(P["passacc"]) in ["nadprum","elite"])))
    mid   = (bucket(P["defd"]) in ["nadprum","elite"] or bucket(P["aerial"]) in ["nadprum","elite"] or
             bucket(P["cross"]) in ["nadprum","elite"] or bucket(P["drib"]) in ["nadprum","elite"])
    if tone == "Přísný":
        if not upper and not mid:
            return "Nedoporučuji do klubů z horní poloviny tabulky; vhodný maximálně jako šířka kádru pro střed/dolní část."
        if upper:
            return "Vhodný pro ambiciózní horní polovinu tabulky; do dominantního prostředí pouze s jasnou rolí a ochranou slabin."
        return "Použitelný pro střed tabulky; do top projektů pouze pod konkrétní herní plán, jinak nedoporučuji."
    else:
        return "Reálně využitelný pro střed až horní střed tabulky; do topu po potvrzení konzistence ve finální třetině."

# --- Narativní report ---
if report_style.startswith("Narativní"):
    # archetyp a úvod
    tags, arche = archetype_sentence(primary_pos, P)
    uvod = (
        f"{jmeno} ({int(vek) if not pd.isna(vek) else 'věk ?'}, {klub}) je typologicky {tags if tags else 'univerzální profil'}. "
        f"{arche} V této sezóně odehrál {int(minuty) if not pd.isna(minuty) else '?'} minut; "
        f"porovnáváno se vzorkem stejné pozice ({primary_pos}, N={n_ref})."
    )

    # herní chování (s míčem i bez) – plynulý text
    attack_bits = []
    if bucket(P["drib"]) in ["nadprum","elite"]:
        attack_bits.append("v 1v1 je průrazný a umí měnit rytmus")
    if bucket(P["cross"]) in ["nadprum","elite"]:
        attack_bits.append("doručuje kvalitní míče z kraje/halfspace")
    if bucket(P["keyp90"]) in ["nadprum","elite"]:
        attack_bits.append("má nadprůměrnou poslední přihrávku")
    if bucket(P["g90"]) in ["nadprum","elite"]:
        attack_bits.append("má opakovatelný gólový výstup")
    if not attack_bits:
        attack_bits.append("s míčem volí bezpečné, funkční řešení")

    defend_bits = []
    if bucket(P["defd"]) in ["nadprum","elite"]:
        defend_bits.append("1v1 zvládá s dobrým načasováním a kontaktem")
    if bucket(P["aerial"]) in ["nadprum","elite"]:
        defend_bits.append("je spolehlivý ve vzduchu a kryje zadní tyč")
    if not defend_bits:
        defend_bits.append("defenzivně spoléhá spíše na poziční hru")

    profil = (
        " Herně s míčem " + (", ".join(attack_bits)) + "; "
        "bez míče " + (", ".join(defend_bits)) + "."
    )

    # limity/rizika – krátké, konkrétní
    risks = []
    if bucket(P["aerial"]) in ["slabe","podprum"] and primary_pos in ["CB","RB","LB","RWB","LWB"]:
        risks.append("riziko na zadní tyči a při standardkách")
    if bucket(P["keyp90"]) in ["slabe","podprum"] and primary_pos in ["LW","RW","AMF"]:
        risks.append("nižší kvalita poslední přihrávky pod tlakem")
    if bucket(P["g90"]) in ["slabe","podprum"] and primary_pos in ["CF","LW","RW"]:
        risks.append("neefektivní koncovka vzhledem k objemu")
    limity = " Rizika: " + ", ".join(risks) + "." if risks else ""

    # přenositelnost – kam sedne
    fit = []
    if primary_pos in ["CB"]:
        fit.append("sedí do trojice stoperů se zajištěním prostoru")
        if bucket(P["passacc"]) in ["nadprum","elite"]:
            fit.append("v dvojici obstojí, pokud má po boku rychlejšího partnera")
    elif primary_pos in ["RB","LB","RWB","LWB"]:
        fit.append("lepší v týmu s přechodem a vysokým postavením krajů")
    elif primary_pos in ["LW","RW","LWF","RWF"]:
        fit.append("nejvíc vytěží z izolací 1v1 a z rychlých přechodů")
    elif primary_pos in ["CF","ST","CF9"]:
        fit.append("uplatní se v boxových vzorcích a řízeném presinku")
    else:
        fit.append("role dle herního plánu, důležitá je kompaktnost mezi liniemi")
    fit_sentence = " Herní využití: " + ", ".join(fit) + "."

    # doporučení (přísné/konzervativní)
    verdict = " Doporučení: " + strict_recommendation(primary_pos, P, tone)

    # doplňující kvant: (jen 2–3 čísla, v závorkách)
    kvant = (
        " (kontext: přesnost přihrávek " + fmt_percentil(P["passacc"]) +
        ", vzdušné souboje " + fmt_percentil(P["aerial"]) +
        ", key passes " + fmt_percentil(P["keyp90"]) + ")."
    )

    narrative = uvod + profil + limity + fit_sentence + verdict + kvant
    st.write(narrative)

else:
    # fallback na tvůj původní strukturovaný text – použij, pokud chceš zachovat obě varianty
    detail_level = st.selectbox("Úroveň detailu reportu", ["Stručný", "Standard", "Obsáhlý"], index=2)

    # (sem můžeš ponechat tvůj existující strukturovaný generátor z předchozí verze;
    # pro zkrácení odpovědi ho sem znovu nekopíruju – v appce jej už máš)
    st.info("Strukturovaný mód je dostupný v předchozí verzi bloku. Doporučuji použít Narativní typologii.")


    # Standardky a variabilita
    std_notes = []
    if bucket(P["aerial"]) in ["nadprum", "elite"]:
        std_notes += ["útočné standardky – náběh na zadní prostor"]
    if primary_pos in ["RB", "LB", "RWB", "LWB"] and bucket(P["cross"]) in ["nadprum", "elite"]:
        std_notes += ["rohy/volné kopy z křídel (doručování)"]
    paragraphs.append("**Standardky.** " + (", ".join(std_notes) if std_notes else "nevýrazný vliv."))

    paragraphs.append("**Variabilita/role.** Může alternovat v rámci postu podle herního plánu; posun do jiné role podmíněn zachováním produkce ve finále.")

    # Rizika
    risks = []
    if bucket(P["aerial"]) in ["slabe", "podprum"] and primary_pos in ["CB", "FB", "RB", "LB", "RWB", "LWB"]:
        risks.append("vzdušné situace na zadní tyči")
    if bucket(P["keyp90"]) in ["slabe", "podprum"] and primary_pos in ["AMF", "LW", "RW"]:
        risks.append("kvalita poslední přihrávky pod tlakem")
    if bucket(P["g90"]) in ["slabe", "podprum"] and primary_pos in ["CF", "LW", "RW"]:
        risks.append("nízká koncovka vůči objemu střel")
    paragraphs.append("**Rizikový profil.** " + (", ".join(risks) + "." if risks else "bez zásadního rizika; sledovat konzistenci výkonu."))

    # Doporučení
    if bucket(P["g90"]) in ["elite", "nadprum"] or bucket(P["keyp90"]) in ["elite", "nadprum"]:
        recomend = "Vhodný pro ambiciózní horní polovinu tabulky; přenos do dominantního prostředí realistický."
    elif bucket(P["drib"]) in ["elite", "nadprum"] and primary_pos in ["LW", "RW", "RWB", "LWB"]:
        recomend = "Smysluplný signing pro týmy s přechodem a 1v1 na krajích; do topu po potvrzení finálního výstupu."
    else:
        recomend = "Použitelný střed ligy; do top pouze při zlepšení ve finále a stabilitě výkonu."
    paragraphs.append("**Doporučení.** " + recomend)

    # Individuální plán
    kpi = [
        "Finále v boxu – automatizovat vzorce (cutback, první dotek, second post).",
        "Tempo poslední přihrávky – volba dřív, kontrola váhy přihrávek do náběhu.",
        "Defenzivní 1v1 – úhly zavírání, práce tělem bez faulu.",
    ]
    paragraphs.append("**Plán 8–12 týdnů.** " + " ".join([f"{i+1}) {t}" for i, t in enumerate(kpi)]))

    # Délka dle volby
    if detail_level == "Stručný":
        st.write("\n\n".join(paragraphs[:3] + paragraphs[-2:]))
    elif detail_level == "Standard":
        st.write("\n\n".join(paragraphs[:6] + paragraphs[-2:]))
    else:
        st.write("\n\n".join(paragraphs))

    # -------------------- Export reportu (DOCX) --------------------
    st.subheader("Export reportu (DOCX)")
    if HAVE_DOCX:
        # Re-render obrázků pro export, aby byly živé instance
        # RADAR
        buf_radar = io.BytesIO()
        vals = np.array([v if not pd.isna(v) else 0 for v in radar_vals])
        vals = np.clip(vals, 0, 150)
        angles = np.linspace(0, 2*np.pi, len(vals), endpoint=False)
        vals_closed = np.concatenate([vals, vals[:1]])
        angles_closed = np.concatenate([angles, angles[:1]])
        fig_r, ax_r = plt.subplots(figsize=(3.8, 3.8), subplot_kw=dict(polar=True))
        ax_r.plot(angles_closed, vals_closed, linewidth=1.6, label=jmeno)
        ax_r.fill(angles_closed, vals_closed, alpha=0.18)
        ax_r.plot(angles_closed, np.ones_like(vals_closed)*100, linestyle="--", linewidth=1)
        ax_r.set_xticks(angles)
        ax_r.set_xticklabels(radar_labels, fontsize=7)
        ax_r.set_yticks([50, 100, 150])
        ax_r.set_yticklabels(["50%", "100%", "150%"], fontsize=7)
        ax_r.set_ylim(0, 150)
        fig_r.savefig(buf_radar, format='png', dpi=200, bbox_inches='tight')
        plt.close(fig_r)
        buf_radar.seek(0)

        # HEATMAP
        buf_heat = io.BytesIO()
        hm = np.array([[v if not pd.isna(v) else np.nan] for v in heat_vals], dtype=float)
        hm_plot = np.nan_to_num(hm, nan=100.0)
        fig_h, ax_h = plt.subplots(figsize=(3.3, 0.26*len(heat_labels)+0.7))
        cmap = LinearSegmentedColormap.from_list("r2g", ["#b30000", "#ff6b6b", "#ffd11a", "#b7e1a1", "#1e7a1e"])
        ax_h.imshow(hm_plot, cmap=cmap, vmin=0, vmax=150)
        ax_h.set_yticks(range(len(heat_labels)))
        ax_h.set_yticklabels(heat_labels, fontsize=8)
        ax_h.set_xticks([])
        for i, v in enumerate(heat_vals):
            txt = "bez dat" if pd.isna(v) else f"{v:.1f}"
            ax_h.text(0, i, txt, ha='center', va='center', fontsize=8, fontweight='bold', color='black')
        fig_h.savefig(buf_heat, format='png', dpi=200, bbox_inches='tight')
        plt.close(fig_h)
        buf_heat.seek(0)

        doc = Document()
        doc.add_heading(f"Typologie hráče – {jmeno}", level=0)
        meta = doc.add_paragraph()
        meta.add_run(f"Klub: {klub} | Pozice: {pozice_raw}\n").bold = True
        meta.add_run(f"Referenční vzorek: stejná pozice {primary_pos}, N={n_ref}\n")
        meta.add_run(f"Minuty: {int(minuty) if not pd.isna(minuty) else '?'} | Věk: {int(vek) if not pd.isna(vek) else '?'}\n")
        doc.add_paragraph(datetime.now().strftime("Vygenerováno: %d.%m.%Y %H:%M"))

        doc.add_heading("Scouting report (souvislý)", level=1)
        for p in paragraphs:
            doc.add_paragraph(p)

        doc.add_heading("Radar (% vs. liga – stejná pozice)", level=1)
        doc.add_picture(buf_radar, width=Inches(3.2))

        doc.add_heading("Heatmapa (0–150 % – stejná pozice)", level=1)
        doc.add_picture(buf_heat, width=Inches(3.0))

        out = io.BytesIO()
        doc.save(out)
        out.seek(0)

        st.download_button(
            label="⬇️ Stáhnout scouting report (.docx)",
            data=out,
            file_name=f"Typologie_{jmeno.replace(' ','_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.warning("DOCX export není dostupný – chybí python-docx. Přidej ho do requirements.txt a redeployni appku.")

else:
    st.info("Nahraj ligový soubor a vyber/nahraj hráče – pak ti vygeneruju vizuály a stažitelný report.")



